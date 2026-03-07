from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 创建文档
doc = Document()

# 设置默认字体为支持中文的字体
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# 添加标题
title = doc.add_heading('智能生活助理应用创意方案', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 副标题
subtitle = doc.add_paragraph('基于 AI 语音助手项目的创意延伸')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(12)
subtitle_format.font.italic = True
subtitle_format.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# 第一部分
doc.add_heading('🚀 应用创意：「智能生活助理 - LifeFlow」', 1)

# 小节 1
doc.add_heading('1️⃣ 应用设想（挖掘想象潜力）', 2)
doc.add_paragraph(
    '在现有 AI 语音助手的基础上，我设想一个多场景智能生活助理：',
    style='Normal'
)
doc.add_paragraph(
    '用户只需按住手机或戴上耳机说一句话，系统就能跨越应用边界智能处理：'
)

# 应用场景列表
scenarios = [
    ('📱', '帮我回复微信说晚饭见', '自动找到微信对话，生成回复'),
    ('🗓️', '周五上午10点提醒我交报告', '自动创建日历提醒'),
    ('🎵', '播放让我放松的音乐', '根据语境识别需求，自动筛选歌单'),
    ('📝', '把今天的想法记录下来', '边说边记，自动整理成日记'),
    ('🚗', '导航到常去的那家咖啡馆', '理解"常去的"，调起地图导航'),
]

for emoji, command, result in scenarios:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{emoji} "{command}"').bold = True
    p.add_run(f' → {result}')

doc.add_paragraph(
    '通过无障碍服务 + LLM 理解，真正做到"一句话管理整个手机"。'
)

doc.add_paragraph()

# 小节 2
doc.add_heading('2️⃣ 具体应用场景（发挥表达能力）', 2)

# 场景一
doc.add_heading('场景一：通勤族的"第二秘书"', 3)
scenario1_text = '''早上7:50，小李开车上班。
他按住耳机说："告诉我今天的日程，天气怎样，还有未读的重要邮件"
系统在5秒内：
  ✓ 从日历拉出3场会议
  ✓ 调用天气服务说"今天多云，温度18-24°C，有50%概率下午下雨"
  ✓ 从邮件客户端筛出3封标记为重要的邮件概览
  
他继续说："提醒我到公司前买杯咖啡"
不需打开任何应用，系统自动在15分钟后发送提醒。'''
doc.add_paragraph(scenario1_text)

# 场景二
doc.add_heading('场景二：夜班医生的"临床笔记官"', 3)
scenario2_text = '''医生查房时，一边观察患者一边说：
"患者李某，血压升高，心率78，建议调整用药方案"
系统实时转录 → AI 格式化 → 自动填入电子病历的标准模板
省去了下班后补记录的2小时。'''
doc.add_paragraph(scenario2_text)

# 场景三
doc.add_heading('场景三：家长的"智能家务助手"', 3)
scenario3_text = '''妈妈做饭时，手上都是油，说："提醒孩子做完作业，
设个45分钟的番茄钟，还有别忘了买牛奶"
系统同时：
  ✓ 给小孩发送提醒（可配置信任的通讯对象）
  ✓ 启动倒计时提醒
  ✓ 在购物清单里加"牛奶"'''
doc.add_paragraph(scenario3_text)

doc.add_paragraph()

# 小节 3
doc.add_heading('3️⃣ 魅力所在（展示雄辩才能）', 2)

# 魅力1
doc.add_heading('其一：破除交互中的"摩擦"', 3)
p1 = doc.add_paragraph()
p1.add_run('传统交互是：').bold = True
p1.add_run('识别意图 → 打开特定应用 → 在应用内完成任务（3-5步）\n')
p1.add_run('这个应用的魅力在于：').bold = True
p1.add_run('识别意图 → 直接完成跨应用任务（1步）')

doc.add_paragraph(
    '对于忙碌的现代人，每一次少打开一个应用，就少一次分心。无论开车、做饭、育儿，双手都被占用的时刻，它成了那个"虚拟秘书"。'
)

doc.add_paragraph()

# 魅力2
doc.add_heading('其二：用 LLM 统一"人机对话语言"', 3)
doc.add_paragraph(
    '现在每个应用都有自己的逻辑和术语。用户要学会：'
)

terminology = [
    '微信的"置顶"和 QQ 的"星标"',
    '日历的"事件"和 todo 的"任务"',
    '导航的"家"和通讯录的"标签"',
]

for term in terminology:
    doc.add_paragraph(term, style='List Bullet')

doc.add_paragraph(
    '而这个应用通过大模型的理解能力，让用户只需说自然语言。同一句话能被翻译成10个不同应用的指令。消除了应用之间的语言壁垒。'
)

doc.add_paragraph()

# 魅力3
doc.add_heading('其三：隐私和自主的平衡', 3)
doc.add_paragraph(
    '与云端 AI(如 Siri、Google Assistant)不同，这个应用可本地化语言模型。用户的每一句话、每一个操作意图，都停留在自己的手机里。安全、隐私、自主性三者统一——这在 AI 时代尤其珍贵。'
)

doc.add_paragraph()

# 魅力4
doc.add_heading('其四：长尾需求的"个性化赋能"', 3)
doc.add_paragraph(
    '主流应用只优化高频需求，冷门但关键的交接需求往往无人问津。比如：'
)

tail_needs = [
    '将手写笔记 + 照片 + 语音一起保存到特定云盘',
    '从多个来源汇总数据生成周报',
    '自动分类和标签多个应用的任务',
]

for need in tail_needs:
    doc.add_paragraph(need, style='List Bullet')

doc.add_paragraph(
    '传统应用永远做不完这些，但一个智能中枢可以。'
)

doc.add_paragraph()

# 总结
doc.add_heading('总结', 2)
summary = doc.add_paragraph()
summary.add_run('这应用的本质魅力是——').font.size = Pt(12)
run = summary.add_run('"从"我去控制手机"变成"手机去理解我""')
run.bold = True
run.font.size = Pt(12)

# 页脚
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
footer_run = footer.add_run('基于 AI_APP 项目的创意思考')
footer_run.font.size = Pt(10)
footer_run.font.italic = True
footer_run.font.color.rgb = RGBColor(150, 150, 150)

# 保存文档
output_path = 'e:/AI_APP/智能生活助理应用方案.docx'
doc.save(output_path)
print(f'✅ 文档已生成：{output_path}')
